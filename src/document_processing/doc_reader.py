"""
Module for reading .doc and .docx files.
Uses python-docx for .docx files and a combination of approaches for legacy .doc files.
"""

import os
import subprocess
from pathlib import Path
from typing import Tuple
import tempfile

from docx import Document
from utils.logger import setup_logger

# Configure logger for this module
logger = setup_logger(__file__)

def extract_doc_text(file_path: str | Path) -> Tuple[str, str]:
    """
    Extract text from .doc or .docx files.
    
    Args:
        file_path: Path to the document file
        
    Returns:
        Tuple[str, str]: (extracted text, document title)
    """
    file_path = Path(file_path)
    filename = file_path.name
    logger.info(f"📄 Extracting text from {filename}")
    
    try:
        if file_path.suffix.lower() == '.docx':
            return _extract_docx(file_path)
        elif file_path.suffix.lower() == '.doc':
            return _extract_doc(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_path.suffix}")
            
    except Exception as e:
        logger.error(f"Error extracting text from {filename}: {str(e)}")
        raise

def _extract_docx(file_path: Path) -> Tuple[str, str]:
    """Extract text from .docx file using python-docx"""
    try:
        doc = Document(file_path)
        
        # Extract title from first paragraph if it looks like a title
        title = None
        if doc.paragraphs:
            first_para = doc.paragraphs[0].text.strip()
            if first_para and (first_para.isupper() or len(first_para.split()) <= 10):
                title = first_para
        
        # If no title found in content, use filename
        if not title:
            title = file_path.stem
            
        # Extract text from paragraphs
        text = []
        for para in doc.paragraphs:
            if para.text.strip():  # Skip empty paragraphs
                text.append(para.text)
                
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text.append(" | ".join(row_text))
                    
        full_text = "\n\n".join(text)
        
        # Add metadata header
        text_with_metadata = f"""Document Metadata:
- Title: {title}
- Filename: {file_path.name}
- Type: DOCX

Content:
{full_text}
"""
        return text_with_metadata, title
        
    except Exception as e:
        logger.error(f"Error processing .docx file {file_path}: {str(e)}")
        raise

def _extract_doc(file_path: Path) -> Tuple[str, str]:
    """
    Extract text from legacy .doc file using multiple fallback methods:
    1. Try converting to DOCX using LibreOffice (if available)
    2. Try using antiword (if available)
    """
    try:
        # First try: Convert to DOCX using LibreOffice
        try:
            with tempfile.TemporaryDirectory() as temp_dir:
                temp_docx = Path(temp_dir) / f"{file_path.stem}.docx"
                
                # Try using LibreOffice to convert with full path
                soffice_path = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
                if not Path(soffice_path).exists():
                    raise FileNotFoundError("LibreOffice not found at expected location")
                    
                result = subprocess.run([
                    soffice_path,
                    '--headless',
                    '--convert-to', 'docx',
                    '--outdir', temp_dir,
                    str(file_path)
                ], check=True, capture_output=True, text=True)
                
                logger.debug(f"LibreOffice conversion output: {result.stdout}")
                if result.stderr:
                    logger.warning(f"LibreOffice conversion warnings: {result.stderr}")
                
                if temp_docx.exists():
                    return _extract_docx(temp_docx)
                else:
                    raise FileNotFoundError(f"Expected output file {temp_docx} not found")
                    
        except (subprocess.SubprocessError, FileNotFoundError) as e:
            logger.warning(f"LibreOffice conversion failed: {str(e)}, trying antiword...")
            
        # Second try: Use antiword
        try:
            result = subprocess.run(
                ['antiword', str(file_path)],
                check=True,
                capture_output=True,
                text=True
            )
            text = result.stdout
            
            if not text.strip():
                raise RuntimeError("antiword produced empty output")
            
            # Clean up the text
            lines = text.split('\n')
            cleaned_lines = []
            title = None
            
            # Process lines and try to extract title
            for line in lines:
                line = line.strip()
                if line:
                    if not title and (line.isupper() or len(line.split()) <= 10):
                        title = line
                    cleaned_lines.append(line)
                    
            # If no title found in content, use filename
            if not title:
                title = file_path.stem
                
            full_text = "\n\n".join(cleaned_lines)
            
            # Add metadata header
            text_with_metadata = f"""Document Metadata:
- Title: {title}
- Filename: {file_path.name}
- Type: DOC
- Extraction Method: antiword

Content:
{full_text}
"""
            return text_with_metadata, title
            
        except (subprocess.SubprocessError, FileNotFoundError, RuntimeError) as e:
            logger.warning(f"antiword failed: {str(e)}")
            raise RuntimeError(
                "Could not extract text from .doc file. Please ensure one of these tools is properly installed:\n"
                "1. LibreOffice (soffice)\n"
                "2. antiword\n"
                f"Last error: {str(e)}"
            )
            
    except Exception as e:
        logger.error(f"Error processing .doc file {file_path}: {str(e)}")
        raise 